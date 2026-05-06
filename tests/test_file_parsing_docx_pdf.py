import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from fastapi.testclient import TestClient

from contract_review_api.main import app

client = TestClient(app)


def _write_blank_pdf(path: Path) -> None:
    # Minimal, valid PDF with one blank page.
    # We generate offsets dynamically so the xref table is correct.
    header = b"%PDF-1.4\n"

    objects: list[bytes] = []
    # 1 0 obj: Catalog
    objects.append(b"1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj\n")
    # 2 0 obj: Pages
    objects.append(b"2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj\n")
    # 3 0 obj: Page
    objects.append(
        b"3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 300] /Contents 4 0 R >>endobj\n"
    )
    # 4 0 obj: Contents (empty stream)
    objects.append(b"4 0 obj<< /Length 0 >>stream\n\nendstream\nendobj\n")

    offsets: list[int] = []
    cursor = len(header)
    for obj in objects:
        offsets.append(cursor)
        cursor += len(obj)

    xref_offset = cursor
    xref = [
        b"xref\n",
        b"0 5\n",
        b"0000000000 65535 f \n",
    ]
    for off in offsets:
        xref.append(f"{off:010d} 00000 n \n".encode("ascii"))
    xref_bytes = b"".join(xref)

    trailer = b"trailer<< /Size 5 /Root 1 0 R >>\nstartxref\n" + str(xref_offset).encode("ascii") + b"\n%%EOF\n"

    path.write_bytes(header + b"".join(objects) + xref_bytes + trailer)


def test_review_from_docx_file_path(tmp_path: Path) -> None:
    from docx import Document

    contract_docx = tmp_path / "contract.docx"
    doc = Document()
    doc.add_paragraph("甲方：北京甲公司")
    doc.add_paragraph("乙方：上海乙公司")
    doc.add_paragraph("项目名称：货物采购")
    doc.add_paragraph("合同类型：买卖合同")
    doc.add_paragraph("自2026年1月1日至2026年12月31日")
    doc.save(str(contract_docx))

    resp = client.post(
        "/reviews",
        json={"file_path": str(contract_docx), "ruleset_ids": ["demo"]},
    )
    assert resp.status_code == 200
    body = resp.json()
    assert body["status"] == "completed"
    assert "final_output" in body


def test_review_from_pdf_file_path_smoke_with_base_text(tmp_path: Path) -> None:
    contract_pdf = tmp_path / "contract.pdf"
    _write_blank_pdf(contract_pdf)

    # Provide base text so the pipeline still has readable content even if pdf text extraction is empty.
    base_text = (
        "甲方：北京甲公司\n"
        "乙方：上海乙公司\n"
        "项目名称：货物采购\n"
        "合同类型：买卖合同\n"
        "自2026年1月1日至2026年12月31日"
    )

    resp = client.post(
        "/reviews",
        json={"text": base_text, "file_path": str(contract_pdf), "ruleset_ids": ["demo"]},
    )
    assert resp.status_code == 200
    assert resp.json()["status"] == "completed"
    assert "final_output" in resp.json()

