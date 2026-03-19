#!/usr/bin/env python3
"""一次性脚本：将原始模板转换为占位符模板。

运行后会在 templates_placeholder/ 目录下生成带 {{占位符}} 的模板文件。
"""

import os
import re
import shutil
from docx import Document

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
ORIGINAL_DIR = os.path.join(SCRIPT_DIR, 'templates')
OUTPUT_DIR = os.path.join(SCRIPT_DIR, 'templates_placeholder')


def replace_in_paragraph(paragraph, replacements):
    """替换段落中的文本（跨 run 处理）。"""
    if not paragraph.runs:
        return
    full_text = ''.join(r.text for r in paragraph.runs)
    new_text = full_text
    for old, new in replacements:
        new_text = new_text.replace(old, new)
    if new_text == full_text:
        return
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ''


def regex_replace_in_paragraph(paragraph, patterns):
    """用正则表达式替换段落文本。"""
    if not paragraph.runs:
        return
    full_text = ''.join(r.text for r in paragraph.runs)
    new_text = full_text
    for pattern, repl in patterns:
        new_text = re.sub(pattern, repl, new_text)
    if new_text == full_text:
        return
    paragraph.runs[0].text = new_text
    for r in paragraph.runs[1:]:
        r.text = ''


def replace_in_doc(doc, replacements=None, regex_patterns=None):
    """在整个文档中执行替换。"""
    for para in doc.paragraphs:
        if replacements:
            replace_in_paragraph(para, replacements)
        if regex_patterns:
            regex_replace_in_paragraph(para, regex_patterns)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replacements:
                        replace_in_paragraph(para, replacements)
                    if regex_patterns:
                        regex_replace_in_paragraph(para, regex_patterns)


def create_wtq_template():
    """创建授权委托书占位符模板。"""
    src = os.path.join(ORIGINAL_DIR, '授权委托书（朱才璋）.docx')
    dst = os.path.join(OUTPUT_DIR, '授权委托书.docx')

    doc = Document(src)

    replacements = [
        ('郭譞', '{{委托人}}'),
        ('王晓博', '{{律师1}}'),
        ('杨卓宇', '{{律师2}}'),
        ('陈文瑞', '{{对方当事人}}'),
        ('买卖合同纠纷', '{{案由}}'),
    ]

    regex_patterns = [
        # 身份证号：后面的空白 → 占位符
        (r'(身份证号：)[\s\t]+', r'\g<1>{{身份证号}}'),
        # 地址：后面的空白 → 占位符
        (r'(地址：)[\s\t]+', r'\g<1>{{地址}}'),
        # 代理权限整行替换
        (r'代理权限为：.*', '代理权限为：{{代理权限}}'),
        # 底部委托人签名行（只匹配纯空白行，避免重复替换）
        (r'^(委托人：)\s*$', r'\g<1>{{委托人}}'),
        # 日期行
        (r'年\s+月\s+日\s*', '{{年}}年{{月}}月{{日}}日'),
    ]

    replace_in_doc(doc, replacements, regex_patterns)
    doc.save(dst)
    print(f"  [完成] 授权委托书.docx")


def create_cth_template():
    """创建民事出庭函占位符模板。"""
    src = os.path.join(ORIGINAL_DIR, '民事出庭函（天驰君泰）.docx')
    dst = os.path.join(OUTPUT_DIR, '民事出庭函.docx')

    doc = Document(src)

    regex_patterns = [
        # 函件编号行: [2026]民经字第   号 → 占位符
        (r'\[.*?\].*?字第\s*号', '[{{年份}}]{{案件字号}}字第{{编号}}号'),
        # 法院行
        (r'\s+人民法院：', '{{法院}}人民法院：'),
        # 正文行
        (r'贵院受理的\s+与\s+纠纷一案中，本所接受\s+的委托，现指派\s+律师作为其诉讼代理人。',
         '贵院受理的{{当事人1}}与{{当事人2}}{{纠纷类型}}纠纷一案中，本所接受{{委托人}}的委托，现指派{{指派律师}}律师作为其诉讼代理人。'),
        # 日期行 (中文数字年份)
        (r'二○二[〇一二三四五六七八九]年\s*月\s*日',
         '{{年份中文}}年{{月}}月{{日}}日'),
        # 承办律师和手机
        (r'承办律师：\s+手\s*机：.*',
         '承办律师：{{承办律师}}  手 机：{{手机}}'),
    ]

    replace_in_doc(doc, regex_patterns=regex_patterns)
    doc.save(dst)
    print(f"  [完成] 民事出庭函.docx")


def create_fddbr_template():
    """创建法定代表人身份证明书占位符模板。"""
    src = os.path.join(ORIGINAL_DIR, '法定代表人身份证明书.docx')
    dst = os.path.join(OUTPUT_DIR, '法定代表人身份证明书.docx')

    doc = Document(src)

    replacements = [
        ('刘芳', '{{法定代表人}}'),
        ('执行董事兼总经理', '{{职务}}'),
        ('太原润羽文化传媒有限公司', '{{单位全称}}'),
    ]

    regex_patterns = [
        # 日期行
        (r'年\s+月\s+日', '{{年}}年{{月}}月{{日}}日'),
    ]

    replace_in_doc(doc, replacements, regex_patterns)
    doc.save(dst)
    print(f"  [完成] 法定代表人身份证明书.docx")


def create_wtdlxy_template():
    """创建委托代理协议占位符模板。"""
    src = os.path.join(OUTPUT_DIR, '..', 'templates', '委托代理协议.docx')
    dst = os.path.join(OUTPUT_DIR, '委托代理协议.docx')

    # 使用 textutil 转换的 .docx 文件
    converted = os.path.join(ORIGINAL_DIR, '委托代理协议.docx')
    if not os.path.exists(converted):
        print("  [跳过] 委托代理协议 - 需要先运行 textutil 转换")
        return

    doc = Document(converted)

    replacements = [
        ('岚县君泰矿山工程有限公司', '{{甲方}}'),
        ('李翠君', '{{法定代表人}}'),
        ('山西海云飞文化旅游有限公司、李红丽、程瑞金', '{{对方当事人}}'),
        ('合同纠纷', '{{案由}}'),
        ('王瑞蓉', '{{主办律师}}'),
        ('13509735616', '{{律师电话}}'),
        ('一审、二审、执行', '{{委托事项范围}}'),
        ('小店区人民法院、太原市中级人民法院', '{{审理机关}}'),
    ]

    regex_patterns = [
        # 客户编号
        (r'客户编号\[\s+\]', '客户编号[{{客户编号}}]'),
        # 案件编号
        (r'案件编号\[\s+\]', '案件编号[{{案件编号}}]'),
        # 职务（空白字段）
        (r'(职务：)\s*$', r'\g<1>{{甲方职务}}'),
        # 营业执照号码（空白字段）
        (r'(营业执照号码：)\s*$', r'\g<1>{{营业执照号}}'),
        # 住所（甲方，空白字段）- 只匹配甲方区域
        (r'^(住所：)\s*$', r'\g<1>{{甲方住所}}'),
        # 邮政编码（甲方）
        (r'^(邮政编码：)\s*$', r'\g<1>{{甲方邮编}}'),
        # 联系人
        (r'(联系人：)\s*$', r'\g<1>{{联系人}}'),
        # 联系电话和传真
        (r'(联系电话：)\s+(传真：)\s*', r'\g<1>{{甲方电话}}  \g<2>{{甲方传真}}'),
        # 代理权限
        (r'甲方授权乙方的代理权限是【.*?】', '甲方授权乙方的代理权限是【{{代理权限}}】'),
        # 代理费比例
        (r'金额的\s+\d+\s+%', '金额的  {{代理费比例}}  %'),
        # 签署时间
        (r'签署时间：\d+年\d+月\d+日', '签署时间：{{签署日期}}'),
    ]

    replace_in_doc(doc, replacements, regex_patterns)

    # 单独处理电子信箱字段（甲方和乙方需要区分）
    email_count = 0
    for para in doc.paragraphs:
        full_text = ''.join(r.text for r in para.runs) if para.runs else ''
        if re.match(r'^电子信箱：\s*$', full_text):
            email_count += 1
            if email_count == 1:
                # 甲方邮箱
                para.runs[0].text = '电子信箱：{{甲方邮箱}}'
            else:
                # 乙方（律所）邮箱 - 留空，律所信息固定
                para.runs[0].text = '电子信箱：'
            for r in para.runs[1:]:
                r.text = ''

    doc.save(dst)
    print(f"  [完成] 委托代理协议.docx")


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print("正在创建占位符模板...\n")

    create_wtq_template()
    create_cth_template()
    create_fddbr_template()
    create_wtdlxy_template()

    print(f"\n所有模板已保存到: {OUTPUT_DIR}/")


if __name__ == '__main__':
    main()
